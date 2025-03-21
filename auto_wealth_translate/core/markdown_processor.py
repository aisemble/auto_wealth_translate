"""
Markdown-based document processing module.
"""
import fitz
import re
from typing import List, Dict, Any, Optional
import logging
from pathlib import Path
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import subprocess
import os

logger = logging.getLogger(__name__)

class MarkdownProcessor:
    """Process documents through Markdown format."""
    
    def __init__(self):
        """Initialize the processor."""
        self.md_converter = markdown.Markdown(extensions=['tables', 'fenced_code'])
        
    def pdf_to_markdown(self, pdf_path: str) -> str:
        """
        Convert PDF to Markdown while preserving structure.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            str: Markdown content
        """
        try:
            doc = fitz.open(pdf_path)
            md_content = []
            
            logger.info(f"Converting PDF to Markdown: {pdf_path}")
            logger.info(f"PDF has {len(doc)} pages")
            
            for page_num, page in enumerate(doc):
                # Extract text blocks with their properties
                text_dict = page.get_text("dict")
                blocks = text_dict["blocks"]
                
                logger.info(f"Processing page {page_num+1} with {len(blocks)} blocks")
                
                # Log entire page text for verification
                page_text = page.get_text()
                logger.info(f"Page {page_num+1} raw text sample: {page_text[:200]}...")
                
                for block_idx, block in enumerate(blocks):
                    if "lines" in block:
                        for line_idx, line in enumerate(block["lines"]):
                            line_text = ""
                            for span in line["spans"]:
                                text = span["text"].strip()
                                if not text:
                                    continue
                                    
                                # Log extracted text for debugging
                                logger.debug(f"Block {block_idx}, Line {line_idx}, Text: {text[:50]}{'...' if len(text) > 50 else ''}")
                                line_text += text + " "
                                    
                                # Determine text style
                                font_size = span["size"]
                                font_name = span["font"]
                                is_bold = "bold" in font_name.lower()
                                
                                # Add text with appropriate markdown formatting
                                if font_size > 14:
                                    md_content.append(f"# {text}")
                                elif font_size > 12:
                                    md_content.append(f"## {text}")
                                elif is_bold:
                                    md_content.append(f"**{text}**")
                                else:
                                    md_content.append(text)
                            
                            # Log full line text
                            if line_text.strip():
                                logger.debug(f"Full line text: {line_text.strip()}")
                                    
                # Add page break
                md_content.append("\n---\n")
            
            result = "\n".join(md_content)
            # Print a sample of the extracted content
            logger.info(f"Extracted markdown sample (first 500 chars): {result[:500]}...")
            logger.info(f"Total markdown content length: {len(result)} characters")
            return result
            
        except Exception as e:
            logger.error(f"Error converting PDF to Markdown: {str(e)}")
            raise
            
    def translate_markdown(self, md_content: str, translator) -> str:
        """
        Translate Markdown content while preserving formatting.
        
        Args:
            md_content: Original Markdown content
            translator: Translation service instance
            
        Returns:
            str: Translated Markdown content
        """
        try:
            # Print a sample of the original markdown content
            logger.info(f"Original markdown sample (first 300 chars): {md_content[:300]}...")
            logger.info(f"Translating from {translator.source_lang} to {translator.target_lang}")
            
            # Split content into translatable segments
            segments = self._split_markdown(md_content)
            translated_segments = []
            
            logger.info(f"Split markdown into {len(segments)} segments")
            
            for i, segment in enumerate(segments):
                if segment["type"] == "text":
                    # Skip if segment is empty or whitespace
                    if not segment["content"].strip():
                        translated_segments.append(segment)
                        continue
                    
                    # Log the original text segment
                    logger.info(f"Segment {i} (original): {segment['content'][:100]}{'...' if len(segment['content']) > 100 else ''}")
                    
                    try:
                        # Get the target language code
                        target_lang = translator.target_lang
                        
                        # For Chinese translation, use a custom approach
                        if target_lang == "zh":
                            logger.info(f"Using specialized approach for Chinese translation of segment {i}")
                            
                            # Create explicit instruction for Chinese translation
                            instruction = f"""
                            Translate the following text from {translator.language_names.get(translator.source_lang, translator.source_lang)} to Chinese (Simplified).
                            
                            Rules:
                            1. ONLY respond with the translated text
                            2. Keep all formatting and special characters
                            3. Use appropriate Chinese financial terminology
                            4. Ensure output is in UTF-8 encoded Chinese characters
                            5. Do not add any comments, explanations, or notes
                            
                            Text to translate:
                            {segment['content']}
                            """
                            
                            # Use direct call with temperature = 0 for consistency
                            translated_text = translator._translate_with_openai(
                                text=instruction,
                                target_lang="zh",
                                temperature=0.0
                            )
                            
                            # Verify Chinese characters are present
                            has_chinese = any('\u4e00' <= char <= '\u9fff' for char in translated_text)
                            if not has_chinese:
                                logger.warning(f"No Chinese characters found in translation output for segment {i}! Content: {translated_text[:100]}")
                                
                                # Try one more time with simpler instruction
                                retry_instruction = f"将以下文本翻译成中文(不要添加任何解释,只需给出翻译结果):\n\n{segment['content']}"
                                translated_text = translator._translate_with_openai(
                                    text=retry_instruction, 
                                    target_lang="zh",
                                    temperature=0.0
                                )
                                
                                # Check again
                                has_chinese = any('\u4e00' <= char <= '\u9fff' for char in translated_text)
                                if not has_chinese:
                                    logger.error(f"Second attempt also failed to produce Chinese characters for segment {i}")
                                else:
                                    logger.info(f"Second attempt successfully produced Chinese characters for segment {i}")
                            else:
                                logger.info(f"Chinese characters verified in translation output for segment {i}")
                                
                        else:
                            # For other languages, use standard approach
                            translated_text = translator._translate_with_openai(
                                text=segment["content"],
                                target_lang=target_lang,
                                temperature=0.3
                            )
                        
                        # Log the translated segment
                        logger.info(f"Segment {i} (translated): {translated_text[:100]}{'...' if len(translated_text) > 100 else ''}")
                        
                        # Store the translated segment
                        translated_segments.append({
                            "type": "text",
                            "content": translated_text
                        })
                        
                    except Exception as e:
                        logger.error(f"Error translating segment {i}: {str(e)}")
                        # Fall back to original content on error
                        translated_segments.append(segment)
                else:
                    # Keep formatting intact
                    translated_segments.append(segment)
            
            # Reconstruct markdown
            result = self._reconstruct_markdown(translated_segments)
            
            # Print a sample of the translated content
            logger.info(f"Translated markdown sample (first 300 chars): {result[:300]}...")
            logger.info(f"Total translated content length: {len(result)} characters")
            
            # Check if we have any Chinese characters in result when target is Chinese
            if translator.target_lang == "zh":
                has_chinese = any('\u4e00' <= char <= '\u9fff' for char in result)
                if not has_chinese:
                    logger.error("No Chinese characters found in the final translated content!")
                else:
                    logger.info("Chinese characters verified in the translated content")
            
            return result
            
        except Exception as e:
            logger.error(f"Error translating Markdown: {str(e)}")
            raise
            
    def _split_markdown(self, md_content: str) -> List[Dict[str, Any]]:
        """Split markdown into translatable segments."""
        segments = []
        current_text = []
        
        for line in md_content.split("\n"):
            # Check for headers
            if line.startswith("#"):
                if current_text:
                    segments.append({
                        "type": "text",
                        "content": "\n".join(current_text)
                    })
                    current_text = []
                segments.append({
                    "type": "format",
                    "content": line
                })
                continue
                
            # Check for formatting
            if line.startswith(("**", "*", "`", ">", "-", "1.", "|")):
                if current_text:
                    segments.append({
                        "type": "text",
                        "content": "\n".join(current_text)
                    })
                    current_text = []
                segments.append({
                    "type": "format",
                    "content": line
                })
                continue
                
            # Regular text
            current_text.append(line)
            
        if current_text:
            segments.append({
                "type": "text",
                "content": "\n".join(current_text)
            })
            
        return segments
        
    def _reconstruct_markdown(self, segments: List[Dict[str, Any]]) -> str:
        """Reconstruct markdown from translated segments."""
        return "\n".join(segment["content"] for segment in segments)
        
    def markdown_to_docx(self, md_content: str, output_path: str):
        """
        Convert Markdown to DOCX.
        
        Args:
            md_content: Markdown content
            output_path: Output DOCX path
        """
        try:
            # Convert markdown to HTML
            html = self.md_converter.convert(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            
            # Create new document
            doc = Document()
            
            # Process HTML elements
            for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'li', 'table']):
                if element.name.startswith('h'):
                    level = int(element.name[1])
                    p = doc.add_paragraph()
                    p.style = f'Heading {level}'
                    p.add_run(element.get_text())
                elif element.name == 'p':
                    p = doc.add_paragraph()
                    p.add_run(element.get_text())
                elif element.name in ['ul', 'ol']:
                    for li in element.find_all('li'):
                        p = doc.add_paragraph()
                        p.style = 'List Bullet' if element.name == 'ul' else 'List Number'
                        p.add_run(li.get_text())
                elif element.name == 'table':
                    table = doc.add_table(rows=1, cols=1)
                    for row in element.find_all('tr'):
                        cells = row.find_all(['td', 'th'])
                        if len(cells) > table.columns:
                            table.add_column()
                        row_cells = table.add_row().cells
                        for i, cell in enumerate(cells):
                            row_cells[i].text = cell.get_text()
            
            # Save document
            doc.save(output_path)
            
        except Exception as e:
            logger.error(f"Error converting Markdown to DOCX: {str(e)}")
            raise
            
    def markdown_to_pdf(self, md_content: str, output_path: str):
        """
        Convert Markdown to PDF using WeasyPrint.
        
        Args:
            md_content: Markdown content
            output_path: Output PDF path
        """
        try:
            import weasyprint
            from weasyprint import HTML, CSS
            from weasyprint.text.fonts import FontConfiguration
            
            # Print a sample of the markdown content to be converted to PDF
            logger.info(f"Converting markdown to PDF, content sample: {md_content[:300]}...")
            
            # Check for Chinese characters
            has_chinese = any('\u4e00' <= char <= '\u9fff' for char in md_content)
            if has_chinese:
                logger.info("Chinese characters detected in markdown content!")
            else:
                logger.warning("No Chinese characters found in the markdown content to convert to PDF!")
            
            # Convert markdown to HTML
            html_content = self.md_converter.convert(md_content)
            
            # Add proper HTML structure with CSS for better formatting and CJK support
            font_css = """
                @font-face {
                    font-family: 'Noto Sans CJK';
                    src: local('Noto Sans CJK SC'), local('Microsoft YaHei'), local('SimHei'), local('SimSun');
                }
            """
            
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    @page {{ size: A4; margin: 2cm; }}
                    
                    body {{
                        font-family: 'Noto Sans CJK', Arial, sans-serif;
                        margin: 40px;
                        line-height: 1.6;
                        font-size: 12pt;
                    }}
                    
                    h1 {{ font-size: 24pt; margin-top: 24pt; margin-bottom: 12pt; }}
                    h2 {{ font-size: 18pt; margin-top: 18pt; margin-bottom: 9pt; }}
                    h3 {{ font-size: 14pt; margin-top: 14pt; margin-bottom: 7pt; }}
                    
                    p {{ margin-bottom: 12pt; }}
                    
                    table {{ 
                        border-collapse: collapse; 
                        width: 100%; 
                        margin-bottom: 20px; 
                        page-break-inside: avoid;
                    }}
                    
                    th, td {{ 
                        border: 1px solid #ddd; 
                        padding: 8px; 
                        text-align: left; 
                    }}
                    
                    th {{ background-color: #f2f2f2; }}
                    
                    .chinese {{ font-family: 'Noto Sans CJK', 'Microsoft YaHei', 'SimHei', 'SimSun', sans-serif; }}
                </style>
            </head>
            <body>
                <div class="chinese">
                    {html_content}
                </div>
            </body>
            </html>
            """
            
            # Log HTML sample
            logger.info(f"Generated HTML sample: {full_html[:300]}...")
            
            # Save to a temporary HTML file for debugging
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as f:
                temp_html_path = f.name
                f.write(full_html.encode('utf-8'))
                
            logger.info(f"Saved HTML to temporary file: {temp_html_path}")
            
            try:
                # Configure fonts
                font_config = FontConfiguration()
                
                # Create PDF with WeasyPrint
                logger.info("Creating PDF with WeasyPrint...")
                
                # Create explicit CSS for CJK font support
                css = CSS(string=font_css, font_config=font_config)
                
                # Convert HTML to PDF
                HTML(string=full_html).write_pdf(
                    output_path,
                    stylesheets=[css],
                    font_config=font_config
                )
                
                logger.info(f"Successfully converted Markdown to PDF: {output_path}")
                return
                
            except Exception as weasyprint_error:
                logger.error(f"WeasyPrint error: {str(weasyprint_error)}")
                # Continue to fallback methods
            
            # If WeasyPrint fails, try with pdfkit
            try:
                import pdfkit
                logger.info("Attempting to use pdfkit as fallback...")
                
                # Try to detect wkhtmltopdf path
                wkhtmltopdf_path = None
                import subprocess
                try:
                    wkhtmltopdf_path = subprocess.check_output(['which', 'wkhtmltopdf']).decode().strip()
                    logger.info(f"Found wkhtmltopdf at: {wkhtmltopdf_path}")
                except:
                    logger.warning("Could not find wkhtmltopdf in PATH")
                
                # Configure pdfkit
                options = {
                    'encoding': 'UTF-8',
                    'page-size': 'A4',
                    'margin-top': '2cm',
                    'margin-right': '2cm',
                    'margin-bottom': '2cm',
                    'margin-left': '2cm',
                    'quiet': ''
                }
                
                config = None
                if wkhtmltopdf_path:
                    config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
                
                # Convert HTML to PDF
                pdfkit.from_string(
                    full_html,
                    output_path,
                    options=options,
                    configuration=config
                )
                
                logger.info(f"Successfully converted Markdown to PDF using pdfkit: {output_path}")
                return
                
            except Exception as pdfkit_error:
                logger.error(f"pdfkit error: {str(pdfkit_error)}")
                # Continue to PyMuPDF fallback
            
            # Last resort: Use PyMuPDF as final fallback
            logger.info("Using PyMuPDF as final fallback...")
            doc = fitz.open()
            page = doc.new_page()
            
            # Use TextWriter for better text layout
            text_writer = fitz.TextWriter(page.rect)
            
            # Add a proper font with CJK support
            font_path = None
            # Try to find a CJK font on the system
            possible_fonts = [
                "/System/Library/Fonts/STHeiti Light.ttc",  # macOS
                "/System/Library/Fonts/PingFang.ttc",        # macOS
                "/usr/share/fonts/truetype/arphic/uming.ttc", # Linux
                "C:/Windows/Fonts/msyh.ttc",                 # Windows
                "C:/Windows/Fonts/simsun.ttc"                # Windows
            ]
            
            for f in possible_fonts:
                if os.path.exists(f):
                    font_path = f
                    logger.info(f"Found CJK font: {font_path}")
                    break
            
            # Split markdown into sections
            sections = md_content.split("\n\n")
            y_position = 50  # Starting position
            
            for section in sections:
                if not section.strip():
                    continue
                
                # Check for headings
                fontsize = 11
                is_heading = False
                if section.startswith("#"):
                    heading_level = min(len(section.split(" ")[0]), 3)  # Cap at h3
                    text = " ".join(section.split(" ")[1:])
                    fontsize = 24 - (heading_level * 4)  # h1=20, h2=16, h3=12
                    is_heading = True
                else:
                    text = section
                
                # Skip empty sections
                if not text.strip():
                    continue
                
                # Log text being added to PDF
                logger.debug(f"Adding section to PDF: {text[:50]}{'...' if len(text) > 50 else ''}")
                
                # Use TextWriter to add text with better layout
                rect = fitz.Rect(50, y_position, page.rect.width - 50, y_position + 100)
                
                # Use a CJK-compatible font if available
                font = font_path if font_path else "helv"
                
                # Different formatting for headings
                if is_heading:
                    text_writer.fill_textbox(
                        rect,
                        text,
                        fontname=font,
                        fontsize=fontsize,
                        color=(0, 0, 0.8),  # Blue-black for headings
                        align=fitz.TEXT_ALIGN_LEFT
                    )
                    y_position += fontsize + 10
                else:
                    # Handle paragraphs
                    text_spans = fitz.TextWriter.fill_textbox(
                        text_writer,
                        rect,
                        text,
                        fontname=font,
                        fontsize=fontsize,
                        align=fitz.TEXT_ALIGN_LEFT
                    )
                    # Move position down based on used height
                    if text_spans > 0:
                        # Approximate height used
                        y_position += text_spans * (fontsize + 2)
                    else:
                        # Fallback if we can't determine
                        y_position += fontsize + 10
                
                # Add extra space between sections
                y_position += 5
                
                # Check if we need a new page
                if y_position > page.rect.height - 50:
                    # Write current text to the page
                    text_writer.write_text(page)
                    # Create a new page
                    page = doc.new_page()
                    # Reset writer and position
                    text_writer = fitz.TextWriter(page.rect)
                    y_position = 50
            
            # Write any remaining text
            text_writer.write_text(page)
            
            # Save the PDF
            doc.save(output_path)
            doc.close()
            
            logger.info(f"Created PDF using PyMuPDF fallback: {output_path}")
            
        except Exception as e:
            logger.error(f"All PDF generation methods failed: {str(e)}", exc_info=True)
            raise 